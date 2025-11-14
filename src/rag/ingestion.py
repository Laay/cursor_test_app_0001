from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List

import docx
from pypdf import PdfReader


@dataclass
class Document:
    """Represents a fully loaded document."""

    source_path: Path
    text: str


class DocumentIngestor:
    """Load plain-text, PDF, and DOCX files into memory."""

    SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}

    def __init__(self, docs_dir: str | Path) -> None:
        self.docs_dir = Path(docs_dir)
        if not self.docs_dir.exists():
            raise FileNotFoundError(f"Docs directory '{self.docs_dir}' does not exist.")

    def load(self) -> List[Document]:
        documents: list[Document] = []
        for path in sorted(self.docs_dir.glob("**/*")):
            if path.is_file() and path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                text = self._read_file(path)
                if text.strip():
                    documents.append(Document(source_path=path, text=text.strip()))
        if not documents:
            raise ValueError(
                f"No supported documents found in '{self.docs_dir}'. "
                "Add .txt, .pdf, or .docx files."
            )
        return documents

    def _read_file(self, path: Path) -> str:
        ext = path.suffix.lower()
        if ext == ".txt":
            return path.read_text(encoding="utf-8")
        if ext == ".pdf":
            reader = PdfReader(str(path))
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n".join(pages)
        if ext == ".docx":
            document = docx.Document(str(path))
            return "\n".join(paragraph.text for paragraph in document.paragraphs)
        raise ValueError(f"Unsupported file extension: {ext}")


class WordChunker:
    """Splits text into overlapping windows measured in words."""

    def __init__(self, chunk_size: int = 300, overlap: int = 50) -> None:
        if chunk_size <= 0:
            raise ValueError("chunk_size must be positive.")
        if overlap < 0:
            raise ValueError("overlap must be >= 0.")
        if overlap >= chunk_size:
            raise ValueError("overlap must be smaller than chunk_size.")
        self.chunk_size = chunk_size
        self.overlap = overlap

    def chunk(self, text: str) -> List[str]:
        words = text.split()
        if not words:
            return []

        stride = self.chunk_size - self.overlap
        chunks: list[str] = []
        for start in range(0, len(words), stride):
            window = words[start : start + self.chunk_size]
            if not window:
                break
            chunks.append(self._clean(" ".join(window)))
        return chunks

    @staticmethod
    def _clean(text: str) -> str:
        return " ".join(text.split())


def chunk_documents(documents: Iterable[Document], chunk_size: int = 300, overlap: int = 50) -> List[dict]:
    """Convert documents into chunk dictionaries with metadata."""

    chunker = WordChunker(chunk_size=chunk_size, overlap=overlap)
    chunk_records: list[dict] = []
    for doc in documents:
        chunks = chunker.chunk(doc.text)
        for idx, chunk_text in enumerate(chunks):
            chunk_records.append(
                {
                    "source": str(doc.source_path),
                    "chunk_id": idx,
                    "text": chunk_text,
                }
            )
    return chunk_records
